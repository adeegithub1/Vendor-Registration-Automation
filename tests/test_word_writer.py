"""
Unit tests for modules/word_writer.py (Phase 6: Word Questionnaire Fill Engine).

Run from the project root with:
    python -m unittest tests.test_word_writer -v

Covers every scenario required by Phase 6's testing deliverable:
    - Table questionnaire
    - Paragraph questionnaire (colon-based + blank-line styles)
    - Mixed questionnaire (table + paragraph in the same document)
    - Empty answer
    - Duplicate questions
    - Long (multiline) answer
    - Missing answer (no AnswerObject for a question)
    - Corrupted document
Plus one additional test proving formatting (bold) is actually preserved,
since that is this phase's central requirement.
"""

import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document

from modules.docx_reader import DocxReadError, load_document
from modules.models import AnswerObject, Location, QuestionObject
from modules.word_writer import (
    ParagraphColonInsertionStrategy,
    TableCellInsertionStrategy,
    WordQuestionnaireFiller,
)


def _make_question(id_, question, table, row, cell) -> QuestionObject:
    return QuestionObject(
        id=id_,
        original_question=question,
        intent=question,
        expected_answer_type="Text",
        possible_document="Company Profile",
        category="Company Information",
        priority="Medium",
        location=Location(table=table, row=row, cell=cell),
    )


class TestTableQuestionnaire(unittest.TestCase):
    """Table questionnaire: only the answer cell should change."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for q in ["Vendor Name", "Registered Address"]:
            cells = table.add_row().cells
            cells[0].text = q
            cells[1].text = ""
        self.doc_path = self.tmp_dir / "table.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_answer_written_and_question_untouched(self):
        document = load_document(self.doc_path)
        questions = [
            _make_question(1, "Vendor Name", table=1, row=1, cell=1),
            _make_question(2, "Registered Address", table=1, row=2, cell=1),
        ]
        answers = [
            AnswerObject(question_id=1, answer="ROTO Pumps Limited", confidence=98, source_document="Company Profile.pdf"),
            AnswerObject(question_id=2, answer="123 Industrial Area, Pune", confidence=90, source_document="Company Profile.pdf"),
        ]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertTrue(all(r.status == "success" for r in results))
        table = document.tables[0]
        self.assertEqual(table.rows[0].cells[1].text, "ROTO Pumps Limited")
        self.assertEqual(table.rows[1].cells[1].text, "123 Industrial Area, Pune")
        # Question cells must be completely untouched.
        self.assertEqual(table.rows[0].cells[0].text, "Vendor Name")
        self.assertEqual(table.rows[1].cells[0].text, "Registered Address")


class TestFormattingPreserved(unittest.TestCase):
    """The central requirement: existing character formatting (bold) must survive a fill."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Bank Name"
        answer_paragraph = table.rows[0].cells[1].paragraphs[0]
        bold_run = answer_paragraph.add_run("")
        bold_run.bold = True
        bold_run.font.size = None  # leave default, just prove .bold survives
        self.doc_path = self.tmp_dir / "formatted.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_bold_formatting_survives_fill(self):
        document = load_document(self.doc_path)
        questions = [_make_question(1, "Bank Name", table=1, row=1, cell=1)]
        answers = [AnswerObject(question_id=1, answer="State Bank of India", confidence=95, source_document="Bank Details.xlsx")]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(results[0].status, "success")
        run = document.tables[0].rows[0].cells[1].paragraphs[0].runs[0]
        self.assertEqual(run.text, "State Bank of India")
        self.assertTrue(run.bold)


class TestParagraphQuestionnaire(unittest.TestCase):
    """Paragraph-style: colon-based (same line) and blank-line-below styles."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        doc.add_paragraph("Company Name:")
        doc.add_paragraph("Year of Establishment")
        doc.add_paragraph("_________________")
        self.doc_path = self.tmp_dir / "paragraph.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_colon_based_same_line(self):
        document = load_document(self.doc_path)
        strategy = ParagraphColonInsertionStrategy()
        strategy.prepare(document)
        question = _make_question(1, "Company Name", table=99, row=99, cell=99)  # nonexistent table -> forces fallback
        position = strategy.try_fill(document, question, "ABC Precision Manufacturing")
        self.assertIsNotNone(position)
        self.assertIn("colon-based", position)
        self.assertIn("ABC Precision Manufacturing", document.paragraphs[0].text)

    def test_blank_line_below(self):
        document = load_document(self.doc_path)
        strategy = ParagraphColonInsertionStrategy()
        strategy.prepare(document)
        question = _make_question(2, "Year of Establishment", table=99, row=99, cell=99)
        position = strategy.try_fill(document, question, "2005")
        self.assertIsNotNone(position)
        self.assertIn("blank line", position)
        self.assertEqual(document.paragraphs[2].text.strip(), "2005")

    def test_full_filler_falls_back_to_paragraph_strategy(self):
        """When the table location doesn't resolve, WordQuestionnaireFiller should use the fallback automatically."""
        document = load_document(self.doc_path)
        questions = [_make_question(1, "Company Name", table=1, row=1, cell=1)]  # no table exists in this doc
        answers = [AnswerObject(question_id=1, answer="ABC Precision Manufacturing", confidence=90, source_document="Company Profile.pdf")]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(results[0].status, "success")
        self.assertIn("Paragraph", results[0].word_position)
        self.assertIn("ABC Precision Manufacturing", document.paragraphs[0].text)


class TestMixedQuestionnaire(unittest.TestCase):
    """A document with both a table and paragraph-style questions."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        doc.add_paragraph("Company Name:")
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        cells = table.add_row().cells
        cells[0].text = "Bank Name"
        cells[1].text = ""
        self.doc_path = self.tmp_dir / "mixed.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_table_and_paragraph_questions_both_resolve_correctly(self):
        document = load_document(self.doc_path)
        questions = [
            _make_question(1, "Bank Name", table=1, row=1, cell=1),  # real table location
            _make_question(2, "Company Name", table=5, row=5, cell=5),  # fake -> falls back to paragraph
        ]
        answers = [
            AnswerObject(question_id=1, answer="HDFC Bank", confidence=92, source_document="Bank Details.xlsx"),
            AnswerObject(question_id=2, answer="ABC Precision Manufacturing", confidence=88, source_document="Company Profile.pdf"),
        ]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertTrue(all(r.status == "success" for r in results))
        self.assertIn("Table", [r for r in results if r.question_id == 1][0].word_position)
        self.assertIn("Paragraph", [r for r in results if r.question_id == 2][0].word_position)


class TestEmptyAnswer(unittest.TestCase):
    """An AnswerObject with an empty/whitespace-only answer should be treated as NOT FOUND, not written blank."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "GST Number"
        table.rows[0].cells[1].text = ""
        self.doc_path = self.tmp_dir / "empty_answer.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_whitespace_only_answer_becomes_not_found(self):
        document = load_document(self.doc_path)
        questions = [_make_question(1, "GST Number", table=1, row=1, cell=1)]
        answers = [AnswerObject(question_id=1, answer="   ", confidence=0, source_document="")]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(results[0].status, "success")
        self.assertEqual(results[0].inserted_answer, "NOT FOUND")
        self.assertEqual(document.tables[0].rows[0].cells[1].text, "NOT FOUND")


class TestDuplicateQuestions(unittest.TestCase):
    """Two rows with IDENTICAL question text must each get their own, correct answer."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for _ in range(2):
            cells = table.add_row().cells
            cells[0].text = "Contact Person"
            cells[1].text = ""
        self.doc_path = self.tmp_dir / "duplicate.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_each_occurrence_gets_its_own_answer(self):
        document = load_document(self.doc_path)
        questions = [
            _make_question(1, "Contact Person", table=1, row=1, cell=1),
            _make_question(2, "Contact Person", table=1, row=2, cell=1),
        ]
        answers = [
            AnswerObject(question_id=1, answer="Rajesh Kumar", confidence=90, source_document="Company Profile.pdf"),
            AnswerObject(question_id=2, answer="Priya Sharma", confidence=85, source_document="Organization Chart.pdf"),
        ]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertTrue(all(r.status == "success" for r in results))
        table = document.tables[0]
        self.assertEqual(table.rows[0].cells[1].text, "Rajesh Kumar")
        self.assertEqual(table.rows[1].cells[1].text, "Priya Sharma")
        self.assertNotEqual(table.rows[0].cells[1].text, table.rows[1].cells[1].text)


class TestLongAnswer(unittest.TestCase):
    """Multiline answers should preserve all lines within the same paragraph (soft breaks, not new paragraphs)."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Describe your Quality Control process"
        table.rows[0].cells[1].text = ""
        self.doc_path = self.tmp_dir / "long_answer.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_multiline_answer_all_lines_present(self):
        document = load_document(self.doc_path)
        questions = [_make_question(1, "Describe your Quality Control process", table=1, row=1, cell=1)]
        long_answer = "Step 1: Incoming inspection.\nStep 2: In-process checks.\nStep 3: Final QC before dispatch."
        answers = [AnswerObject(question_id=1, answer=long_answer, confidence=80, source_document="Quality Manual.pdf")]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(results[0].status, "success")
        cell = document.tables[0].rows[0].cells[1]
        # Paragraph count must stay at 1 (soft breaks, not new paragraphs) --
        # this is the layout-preservation guarantee for multiline text.
        self.assertEqual(len(cell.paragraphs), 1)
        paragraph = cell.paragraphs[0]
        full_text = paragraph.text
        self.assertIn("Step 1", full_text)
        self.assertIn("Step 2", full_text)
        self.assertIn("Step 3", full_text)

        # IMPORTANT: cell.text / paragraph.text do NOT render <w:br/> line
        # break elements as "\n" -- so the assertions above would pass even
        # if line breaks were silently missing from the document (this
        # actually happened during development; see word_writer.py's
        # docstring on _write_text_into_paragraph_preserving_format).
        # Verify the real break elements exist in the underlying XML,
        # which is what Word actually renders as separate lines.
        xml = paragraph._p.xml
        break_count = xml.count("<w:br/>") + xml.count("<w:br ")
        self.assertEqual(
            break_count, 2,
            "Expected 2 real <w:br/> line break elements in the XML (3 lines = 2 breaks) "
            "-- text alone is not sufficient proof of multiline formatting.",
        )


class TestMissingAnswer(unittest.TestCase):
    """A question with NO corresponding AnswerObject at all should get NOT FOUND, with a warning logged, not a crash."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "PAN Number"
        table.rows[0].cells[1].text = ""
        self.doc_path = self.tmp_dir / "missing_answer.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_missing_answer_object_yields_not_found(self):
        document = load_document(self.doc_path)
        questions = [_make_question(1, "PAN Number", table=1, row=1, cell=1)]
        answers = []  # deliberately no AnswerObject for question id=1

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(results[0].status, "success")
        self.assertEqual(results[0].inserted_answer, "NOT FOUND")
        self.assertEqual(document.tables[0].rows[0].cells[1].text, "NOT FOUND")


class TestCorruptedDocument(unittest.TestCase):
    """A corrupted/non-DOCX file must raise a clear error, not crash unpredictably."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        self.corrupt_path = self.tmp_dir / "corrupted.docx"
        self.corrupt_path.write_text("this is not a real docx file")

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_corrupted_document_raises_docx_read_error(self):
        filler = WordQuestionnaireFiller()
        questions = [_make_question(1, "Vendor Name", table=1, row=1, cell=1)]
        answers = [AnswerObject(question_id=1, answer="Test", confidence=90, source_document="Test.pdf")]

        with self.assertRaises(DocxReadError):
            filler.fill_and_save(self.corrupt_path, questions, answers, self.tmp_dir / "output.docx")


class TestQuestionLocationOutOfRange(unittest.TestCase):
    """A location pointing past the end of the document (e.g. stale data) must fail that ONE question, not the batch."""

    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Vendor Name"
        table.rows[0].cells[1].text = ""
        self.doc_path = self.tmp_dir / "short.docx"
        doc.save(self.doc_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_out_of_range_location_fails_gracefully_others_still_succeed(self):
        document = load_document(self.doc_path)
        questions = [
            _make_question(1, "Vendor Name", table=1, row=1, cell=1),  # valid
            _make_question(2, "Nonexistent Question", table=1, row=99, cell=1),  # row out of range, no paragraph match either
        ]
        answers = [
            AnswerObject(question_id=1, answer="ROTO Pumps Limited", confidence=95, source_document="Company Profile.pdf"),
            AnswerObject(question_id=2, answer="Some Answer", confidence=80, source_document="Company Profile.pdf"),
        ]

        filler = WordQuestionnaireFiller()
        results = filler.fill(document, questions, answers)

        self.assertEqual(len(results), 2)
        self.assertEqual(results[0].status, "success")
        self.assertEqual(results[1].status, "failed")
        # The batch must still complete -- question 1's answer must be present
        # despite question 2 failing.
        self.assertEqual(document.tables[0].rows[0].cells[1].text, "ROTO Pumps Limited")


if __name__ == "__main__":
    unittest.main()
