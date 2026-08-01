"""
Word Questionnaire Fill Engine.

WHAT THIS MODULE IS RESPONSIBLE FOR
-------------------------------------
Takes the ORIGINAL questionnaire .docx plus the QuestionObjects (Phase 4)
and AnswerObjects (Phase 5) already produced for it, and writes each
answer into the correct answer location in that same document — never
creating a new document from scratch, never touching question text, and
never altering fonts, colors, spacing, tables, headers/footers, or any
other formatting. The output is the SAME document object, saved to a new
path — this is what makes "everything except the inserted answers stays
visually identical" achievable at all: nothing about the document is
regenerated, only specific run text is mutated in place.

REUSES EXISTING MODULES (per this phase's explicit instruction)
---------------------------------------------------------------------
- modules/docx_reader.py: load_document(), get_tables(), get_row_grid_cells()
  — the exact same merge-safe table traversal Phase 2 used to DETECT
  question/answer cells is reused here to LOCATE those same cells again.
  This is not a coincidence: Location.table/row/cell values were recorded
  using this traversal, so re-using it guarantees the coordinates line up.
- modules/models.py: QuestionObject, AnswerObject, FillResult — no new
  models are duplicated here.
- modules/logger.py: the same centralized logger used everywhere else.

STRATEGY PATTERN FOR INSERTION
-----------------------------------
Two strategies, tried in order, mirroring the Strategy/Registry pattern
used throughout this project (Phase 3's processor registry, Phase 4's
ordered classifier rules, Phase 5's provider registry):

  1. TableCellInsertionStrategy (primary) — uses the question's recorded
     Location (table/row/cell) to go directly to the correct cell. This
     is also why duplicate/identical-looking questions are never at risk
     of being filled in the wrong place: each QuestionObject's location
     was recorded independently at detection time, so two questions with
     identical text still have two distinct, correct locations.
  2. ParagraphColonInsertionStrategy (fallback) — for questions whose
     table location doesn't resolve (e.g. table missing/out of range),
     searches document paragraphs for matching question text and fills
     either directly after a colon on the same line, or the next blank/
     underscore-only line.

SCOPE BOUNDARY, STATED HONESTLY
------------------------------------
Phase 2's detector (explicitly not modified in this phase) only ever
produces table-based locations today — it does not detect paragraph,
colon, or blank-line style questions in the first place. So in the
current end-to-end pipeline, every question resolves via strategy 1.
Strategy 2 is fully implemented and independently unit-tested (see
tests/test_word_writer.py), ready for when non-table detection exists,
but it is not reachable through today's --read-docx / --fill-questionnaire
flow. This is called out explicitly rather than left to be discovered.

HOW FORMATTING IS ACTUALLY PRESERVED
-----------------------------------------
python-docx stores character formatting (font, size, color, bold, etc.)
on individual Run objects, not on the Cell or Paragraph. This module NEVER
deletes or recreates a paragraph or cell — it only ever mutates an
EXISTING run's `.text` attribute, which changes the text while leaving
every formatting property on that run untouched. New runs are only
created when a paragraph had zero runs to begin with (a genuinely empty
cell/line) — in that one case, the new run necessarily takes the
paragraph's default style, since there was no prior run formatting to
preserve. This is the one honest limit of "preserve formatting": you
cannot preserve character formatting that was never there.

MULTILINE ANSWERS
--------------------
Long answers are inserted as multiple lines WITHIN THE SAME PARAGRAPH,
using soft line breaks (a run's add_break()), not new paragraphs. Using
new paragraphs would change the cell's paragraph count and could alter
row height/spacing in ways a reviewer would notice; a soft line break
achieves multi-line text while keeping the paragraph structure (and
therefore layout) untouched.

ERROR HANDLING
---------------
Every question is attempted independently. If writing one question's
answer fails for any reason (location out of range, unexpected document
structure, etc.), it's logged as a warning and processing continues —
per this phase's explicit requirement, one failure never stops the batch.
"""

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional

from docx.document import Document as DocumentObject
from docx.table import Table

from modules.docx_reader import get_row_grid_cells, get_tables, load_document
from modules.logger import get_logger
from modules.models import AnswerObject, FillResult, QuestionObject

logger = get_logger(__name__)

_NOT_FOUND_ANSWER = "NOT FOUND"
_BLANK_LINE_PATTERN = re.compile(r"^_{3,}$")


def _write_text_into_paragraph_preserving_format(paragraph, text: str) -> None:
    """
    Write (possibly multiline) text into an existing paragraph, reusing
    its first run's formatting wherever a run already exists, and using
    soft line breaks (not new paragraphs) for multiline text.

    IMPORTANT: after calling run.add_break(), text for that same run must
    be added with run.add_text(...), NOT by assigning run.text = ... .
    python-docx's Run.text SETTER rebuilds the run's XML children from
    scratch (it does not append), which would silently delete the break
    element just added. add_text() appends a new <w:t> child without
    disturbing any existing children, which is what actually preserves
    the break. This was caught by inspecting the saved document's raw
    XML during testing -- the break elements were silently missing
    despite the code "looking" correct.
    """
    lines = text.split("\n") if text else [""]
    existing_runs = list(paragraph.runs)

    if existing_runs:
        base_run = existing_runs[0]
        base_run.text = lines[0]
        for extra_run in existing_runs[1:]:
            extra_run.text = ""
        for line in lines[1:]:
            new_run = paragraph.add_run()
            new_run.add_break()
            new_run.add_text(line)
    else:
        paragraph.add_run(lines[0])
        for line in lines[1:]:
            new_run = paragraph.add_run()
            new_run.add_break()
            new_run.add_text(line)


class AnswerInsertionStrategy(ABC):
    """A way of locating and filling one question's answer within the document."""

    def prepare(self, document: DocumentObject) -> None:
        """
        Optional one-time setup before processing all questions in a
        document (e.g. caching document.tables so repeated lookups don't
        re-walk the document). Default is a no-op; strategies override
        only if they benefit.
        """
        return None

    @abstractmethod
    def try_fill(self, document: DocumentObject, question_object: QuestionObject, answer_text: str) -> Optional[str]:
        """
        Attempt to locate and fill this question's answer.

        Returns:
            A short human-readable description of where the answer was
            written (e.g. "Table 1, Row 3, Cell 1"), or None if this
            strategy could not find an applicable location (the caller
            will try the next strategy).
        """
        raise NotImplementedError


class TableCellInsertionStrategy(AnswerInsertionStrategy):
    """
    Primary strategy: uses QuestionObject.location (table/row/cell) to go
    directly to the correct answer cell, reusing Phase 2's merge-safe
    table traversal so the coordinates line up exactly with detection.
    """

    def __init__(self):
        self._tables_cache: Optional[List[Table]] = None

    def prepare(self, document: DocumentObject) -> None:
        # Cache once per document instead of re-walking document.tables
        # for every single question -- meaningful on large questionnaires.
        self._tables_cache = get_tables(document)

    def try_fill(self, document: DocumentObject, question_object: QuestionObject, answer_text: str) -> Optional[str]:
        tables = self._tables_cache if self._tables_cache is not None else get_tables(document)
        location = question_object.location

        table_index = location.table - 1
        if table_index < 0 or table_index >= len(tables):
            return None
        table = tables[table_index]

        row_index = location.row - 1
        if row_index < 0 or row_index >= len(table.rows):
            return None
        row = table.rows[row_index]

        grid_cells = get_row_grid_cells(row)
        matching_cells = [gc for gc in grid_cells if gc.column_index == location.cell]
        if not matching_cells:
            return None

        target_cell = matching_cells[0].cell
        paragraph = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
        _write_text_into_paragraph_preserving_format(paragraph, answer_text)

        return f"Table {location.table}, Row {location.row}, Cell {location.cell}"


class ParagraphColonInsertionStrategy(AnswerInsertionStrategy):
    """
    Fallback strategy for non-table questionnaire styles: colon-based
    ("Question:") and question-on-one-line/blank-or-underscore-line-below
    styles. Matches by question text (normalized whitespace/case), then
    fills either right after the colon or the following blank line.
    """

    def __init__(self):
        self._paragraphs_cache = None

    def prepare(self, document: DocumentObject) -> None:
        self._paragraphs_cache = list(document.paragraphs)

    @staticmethod
    def _normalize(text: str) -> str:
        return re.sub(r"\s+", " ", text.strip().lower())

    @staticmethod
    def _is_blank_line(text: str) -> bool:
        stripped = text.strip()
        return stripped == "" or bool(_BLANK_LINE_PATTERN.match(stripped))

    def try_fill(self, document: DocumentObject, question_object: QuestionObject, answer_text: str) -> Optional[str]:
        paragraphs = self._paragraphs_cache if self._paragraphs_cache is not None else list(document.paragraphs)
        target_normalized = self._normalize(question_object.original_question)
        if not target_normalized:
            return None

        for index, paragraph in enumerate(paragraphs):
            paragraph_normalized = self._normalize(paragraph.text)
            if not paragraph_normalized:
                continue
            if target_normalized not in paragraph_normalized and paragraph_normalized not in target_normalized:
                continue

            if ":" in paragraph.text:
                _, _, suffix = paragraph.text.partition(":")
                if suffix.strip() == "":
                    if paragraph.runs:
                        last_run = paragraph.runs[-1]
                        last_run.text = last_run.text.rstrip() + " " + answer_text
                    else:
                        paragraph.add_run(" " + answer_text)
                    return f"Paragraph {index} (colon-based, same line)"

            if index + 1 < len(paragraphs):
                next_paragraph = paragraphs[index + 1]
                if self._is_blank_line(next_paragraph.text):
                    _write_text_into_paragraph_preserving_format(next_paragraph, answer_text)
                    return f"Paragraph {index + 1} (blank line following question)"

        return None


class WordQuestionnaireFiller:
    """Orchestrates filling every question's answer into a questionnaire document."""

    def __init__(self, strategies: Optional[List[AnswerInsertionStrategy]] = None):
        """
        Args:
            strategies: Ordered list of insertion strategies to try, first
                match wins. Defaults to [TableCellInsertionStrategy(),
                ParagraphColonInsertionStrategy()]. Injectable for testing
                or to add/reorder strategies without touching this class.
        """
        self.strategies = strategies if strategies is not None else [
            TableCellInsertionStrategy(),
            ParagraphColonInsertionStrategy(),
        ]

    @staticmethod
    def _resolve_answer_text(question_object: QuestionObject, answer_by_id: Dict[int, AnswerObject]) -> str:
        answer_object = answer_by_id.get(question_object.id)
        if answer_object is None:
            logger.warning(f"No AnswerObject found for question id={question_object.id}; using NOT FOUND.")
            return _NOT_FOUND_ANSWER
        if not answer_object.answer or not answer_object.answer.strip():
            return _NOT_FOUND_ANSWER
        return answer_object.answer

    def fill(
        self,
        document: DocumentObject,
        question_objects: List[QuestionObject],
        answers: List[AnswerObject],
    ) -> List[FillResult]:
        """
        Fill every question's answer into the given (already-loaded)
        document, in place. Does not save the document — call
        fill_and_save() for the full load-fill-save pipeline.

        Returns:
            List of FillResult, one per question, in the same order as
            question_objects. Never raises for an individual question's
            failure — see this module's ERROR HANDLING docstring section.
        """
        answer_by_id = {answer.question_id: answer for answer in answers}

        for strategy in self.strategies:
            strategy.prepare(document)

        results: List[FillResult] = []

        for question_object in question_objects:
            answer_text = self._resolve_answer_text(question_object, answer_by_id)
            question_preview = question_object.original_question[:80]

            try:
                word_position = None
                for strategy in self.strategies:
                    word_position = strategy.try_fill(document, question_object, answer_text)
                    if word_position is not None:
                        break
            except Exception as exc:
                logger.warning(
                    f"Question ID={question_object.id} | Question='{question_preview}' | "
                    f"Answer='{answer_text[:80]}' | Position=N/A | Status=FAILED | Reason={exc}"
                )
                results.append(
                    FillResult(
                        question_id=question_object.id,
                        question_text=question_object.original_question,
                        inserted_answer=answer_text,
                        word_position="",
                        status="failed",
                        error_message=str(exc),
                    )
                )
                continue

            if word_position is None:
                reason = "No matching location found in the document for this question."
                logger.warning(
                    f"Question ID={question_object.id} | Question='{question_preview}' | "
                    f"Answer='{answer_text[:80]}' | Position=N/A | Status=FAILED | Reason={reason}"
                )
                results.append(
                    FillResult(
                        question_id=question_object.id,
                        question_text=question_object.original_question,
                        inserted_answer=answer_text,
                        word_position="",
                        status="failed",
                        error_message=reason,
                    )
                )
                continue

            logger.info(
                f"Question ID={question_object.id} | Question='{question_preview}' | "
                f"Answer='{answer_text[:80]}' | Position={word_position} | Status=SUCCESS"
            )
            results.append(
                FillResult(
                    question_id=question_object.id,
                    question_text=question_object.original_question,
                    inserted_answer=answer_text,
                    word_position=word_position,
                    status="success",
                )
            )

        return results

    def fill_and_save(
        self,
        original_questionnaire_path: Path,
        question_objects: List[QuestionObject],
        answers: List[AnswerObject],
        output_path: Path,
    ) -> List[FillResult]:
        """
        Full pipeline: load the original questionnaire, fill every
        answer, and save the result to a new path (the original file is
        never modified).

        Raises:
            DocxReadError: If the original questionnaire cannot be opened
                (missing, wrong extension, or corrupted). This is a
                whole-document failure, not a per-question one, so it is
                intentionally NOT caught here — the caller decides how to
                handle it (see app.py's run_fill_questionnaire()).
        """
        document = load_document(original_questionnaire_path)
        results = self.fill(document, question_objects, answers)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        logger.info(f"Saved filled questionnaire to {output_path}")

        return results
